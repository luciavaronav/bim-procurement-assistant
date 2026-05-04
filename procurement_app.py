"""
BIM Procurement Assistant — Streamlit App v2
TFM Tool: BIM Schedule → Procurement Pipeline + AI Strategy
Enhancements: Lead Time Estimator, Delivery Sequencing, Spend Estimation
Run with: streamlit run procurement_app_v2.py
"""

import streamlit as st
import pandas as pd
import io
import zipfile
from datetime import datetime, timedelta

# ── PAGE CONFIG ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="BIM Procurement Assistant",
    page_icon="⚡",
    layout="wide",
)

# ── CUSTOM CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .stApp { background-color: #f0f2f6; }
    .metric-card {
        background: white; border-radius: 12px; padding: 20px;
        text-align: center; box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    .metric-number { font-size: 2.5rem; font-weight: 700; margin: 0; }
    .metric-label { color: #666; font-size: 0.9rem; margin: 0; }
    .ai-box {
        background: #fffdf0; border-left: 5px solid #f1c40f;
        padding: 20px; border-radius: 8px;
        white-space: pre-wrap; font-family: monospace; font-size: 0.9rem;
    }
    .section-header {
        background: white; border-radius: 12px; padding: 15px 20px;
        margin: 10px 0; box-shadow: 0 2px 8px rgba(0,0,0,0.06);
    }
</style>
""", unsafe_allow_html=True)

# ── CONSTANTS ─────────────────────────────────────────────────────────────────
EXPECTED_COLS = ["Family", "Type", "Panel Name", "Level", "Count"]

# ── LEAD TIME DATABASE (weeks) ────────────────────────────────────────────────
# Industry-standard lead times by keyword match in family name
# Sources: RS Means, NECA Manual of Labor Units, typical MEP procurement benchmarks
LEAD_TIME_DB = {
    "transformer":          {"min": 16, "max": 24, "unit_cost": 18000},
    "switchboard":          {"min": 14, "max": 20, "unit_cost": 35000},
    "switchgear":           {"min": 20, "max": 30, "unit_cost": 55000},
    "panelboard":           {"min": 8,  "max": 14, "unit_cost": 3500},
    "panel":                {"min": 8,  "max": 14, "unit_cost": 3500},
    "disconnect":           {"min": 4,  "max": 8,  "unit_cost": 800},
    "meter":                {"min": 6,  "max": 10, "unit_cost": 1200},
    "pull box":             {"min": 2,  "max": 4,  "unit_cost": 350},
    "photovoltaic":         {"min": 12, "max": 18, "unit_cost": 4500},
    "pv":                   {"min": 12, "max": 18, "unit_cost": 4500},
    "inverter":             {"min": 10, "max": 16, "unit_cost": 6000},
    "battery":              {"min": 14, "max": 20, "unit_cost": 22000},
    "ups":                  {"min": 8,  "max": 14, "unit_cost": 8000},
    "generator":            {"min": 16, "max": 26, "unit_cost": 45000},
    "ats":                  {"min": 10, "max": 16, "unit_cost": 5500},
    "unbalanced power":     {"min": 6,  "max": 10, "unit_cost": 1800},
    "connector":            {"min": 3,  "max": 6,  "unit_cost": 250},
    "air handler":          {"min": 12, "max": 20, "unit_cost": 28000},
    "vav":                  {"min": 6,  "max": 10, "unit_cost": 1400},
    "fan":                  {"min": 6,  "max": 12, "unit_cost": 3200},
    "chiller":              {"min": 20, "max": 30, "unit_cost": 85000},
    "boiler":               {"min": 16, "max": 24, "unit_cost": 42000},
    "pump":                 {"min": 8,  "max": 14, "unit_cost": 4500},
    "cooling tower":        {"min": 14, "max": 22, "unit_cost": 38000},
    "default":              {"min": 6,  "max": 12, "unit_cost": 1500},
}

def get_lead_time(family_name: str) -> dict:
    """Match family name to lead time entry using keyword lookup."""
    name_lower = family_name.lower()
    for keyword, data in LEAD_TIME_DB.items():
        if keyword in name_lower:
            return {**data, "keyword": keyword}
    return {**LEAD_TIME_DB["default"], "keyword": "default"}


# ── HELPERS ───────────────────────────────────────────────────────────────────

def clean_bim_csv(uploaded_file):
    """Full cleaning pipeline, returns (df_clean, df_family, df_issues, total_units)."""
    df_raw = pd.read_csv(uploaded_file, header=None, names=range(5),
                         dtype=str, sep=',', encoding='utf-8')

    # Detect header row
    header_row_idx = None
    for i in range(min(20, len(df_raw))):
        row_vals = [str(v).strip().lower() for v in df_raw.iloc[i].fillna("")]
        if "family" in row_vals[0]:
            header_row_idx = i
            break

    if header_row_idx is None:
        st.error("❌ Could not find expected header row ('Family') in first 20 rows.")
        st.stop()

    df = df_raw.iloc[header_row_idx + 1:].copy()
    df.columns = EXPECTED_COLS

    for col in EXPECTED_COLS:
        df[col] = df[col].fillna("").astype(str).str.strip()

    df["Count"] = pd.to_numeric(df["Count"], errors='coerce')
    df_clean = df[df["Count"].notna()].copy()
    df_clean["Count"] = df_clean["Count"].astype(int)

    # Remove subtotals / empty families
    mask_subtotal = (
        (df_clean["Type"] == "") &
        (df_clean["Panel Name"] == "") &
        (df_clean["Level"] == "")
    )
    df_clean = df_clean[~mask_subtotal].copy()
    df_clean = df_clean[df_clean["Family"] != ""].copy()

    # Issues
    def get_issues(row):
        errors = []
        if not row["Type"]:       errors.append("Type undefined")
        if not row["Panel Name"]: errors.append("Panel not assigned")
        if not row["Level"]:      errors.append("Level missing")
        return "; ".join(errors)

    df_clean["Issues"] = df_clean.apply(get_issues, axis=1)
    df_issues = df_clean[df_clean["Issues"] != ""].copy()
    df_family = (df_clean.groupby("Family")["Count"].sum()
                 .reset_index().sort_values("Count", ascending=False))
    total_units = int(df_clean["Count"].sum())

    return df_clean, df_family, df_issues, total_units


# ── NEW: LEAD TIME ESTIMATOR ──────────────────────────────────────────────────

def build_lead_time_df(df_family: pd.DataFrame, project_start: datetime) -> pd.DataFrame:
    """
    For each equipment family, calculate:
    - Lead time range (weeks)
    - Order-by date (project_start - max lead time)
    - Status: CRITICAL / WARNING / OK
    """
    rows = []
    for _, row in df_family.iterrows():
        lt = get_lead_time(row["Family"])
        order_by = project_start - timedelta(weeks=lt["max"])
        days_until = (order_by - datetime.now()).days
        if days_until < 0:
            status = "🔴 OVERDUE"
        elif days_until < 14:
            status = "🟠 CRITICAL"
        elif days_until < 30:
            status = "🟡 WARNING"
        else:
            status = "🟢 OK"
        rows.append({
            "Family": row["Family"],
            "Count": int(row["Count"]),
            "Lead Time Min (wks)": lt["min"],
            "Lead Time Max (wks)": lt["max"],
            "Order By Date": order_by.strftime("%Y-%m-%d"),
            "Days Until Order Deadline": days_until,
            "Status": status,
        })
    df = pd.DataFrame(rows).sort_values("Days Until Order Deadline")
    return df


# ── NEW: DELIVERY SEQUENCING ──────────────────────────────────────────────────

def build_delivery_sequence(df_clean: pd.DataFrame) -> pd.DataFrame:
    """
    Group BOM by Level and generate a recommended delivery sequence.
    Priority: Basement/Parking first (main distribution), then floor ascending,
    then Roof last (PV/special equipment).
    """
    level_totals = (df_clean.groupby("Level")["Count"]
                    .sum().reset_index()
                    .rename(columns={"Count": "Total Units"}))

    def level_sort_key(level: str) -> int:
        lvl = level.lower()
        if any(x in lvl for x in ["park", "basement", "b1", "b2", "ground"]):
            return 0
        if "roof" in lvl or "r1" in lvl or "r2" in lvl:
            return 99
        # Extract first number found
        import re
        nums = re.findall(r'\d+', lvl)
        return int(nums[0]) if nums else 50

    level_totals["Sort"] = level_totals["Level"].apply(level_sort_key)
    level_totals = level_totals.sort_values("Sort").drop(columns="Sort").reset_index(drop=True)
    level_totals.index += 1
    level_totals.index.name = "Delivery Phase"

    level_totals["Recommended Action"] = level_totals["Level"].apply(lambda lvl: (
        "Main distribution equipment — schedule first delivery"
        if any(x in lvl.lower() for x in ["park", "basement", "b1", "ground"])
        else "Rooftop/PV equipment — schedule final delivery, coordinate crane access"
        if any(x in lvl.lower() for x in ["roof", "r1", "r2"])
        else "Standard floor delivery — coordinate with installation sequence"
    ))

    return level_totals


# ── NEW: SPEND ESTIMATION ─────────────────────────────────────────────────────

def build_spend_estimation(df_family: pd.DataFrame) -> pd.DataFrame:
    """
    Estimate procurement spend per family using unit cost database.
    Returns df with unit cost, total estimated cost, and spend tier.
    """
    rows = []
    for _, row in df_family.iterrows():
        lt = get_lead_time(row["Family"])
        unit_cost = lt["unit_cost"]
        total_cost = unit_cost * int(row["Count"])
        if total_cost > 100000:
            tier = "🔴 High Value (>$100k)"
        elif total_cost > 20000:
            tier = "🟡 Medium Value ($20k–$100k)"
        else:
            tier = "🟢 Low Value (<$20k)"
        rows.append({
            "Family": row["Family"],
            "Count": int(row["Count"]),
            "Unit Cost (USD)": unit_cost,
            "Estimated Total (USD)": total_cost,
            "Spend Tier": tier,
        })
    df = pd.DataFrame(rows).sort_values("Estimated Total (USD)", ascending=False)
    df["Unit Cost (USD)"] = df["Unit Cost (USD)"].apply(lambda x: f"${x:,.0f}")
    df["Estimated Total (USD)"] = df["Estimated Total (USD)"].apply(lambda x: f"${x:,.0f}")
    return df


def get_total_spend(df_family: pd.DataFrame) -> int:
    total = 0
    for _, row in df_family.iterrows():
        lt = get_lead_time(row["Family"])
        total += lt["unit_cost"] * int(row["Count"])
    return total


# ── AI PROMPT (enhanced with new analytics) ───────────────────────────────────

def build_prompt(total_units, n_rows, n_issues, df_family,
                 df_lead, df_sequence, total_spend) -> str:
    family_data = df_family.to_string(index=False)

    # Critical items (overdue or critical lead time)
    critical = df_lead[df_lead["Status"].str.contains("OVERDUE|CRITICAL", na=False)]
    critical_str = critical[["Family", "Lead Time Max (wks)", "Order By Date"]].to_string(index=False) \
        if len(critical) > 0 else "None identified."

    sequence_str = df_sequence[["Level", "Total Units", "Recommended Action"]].to_string(index=False)

    return f"""You are an office-based Senior Procurement Engineer on an MEP/Electrical project.
You have received a validated BOM exported from Autodesk Revit via a Python pipeline.

VALIDATED DATA - do not recalculate, use exactly as provided:
- Total units: {total_units}
- BOM line items: {n_rows}
- Data quality issues: {n_issues}
- Estimated total procurement spend: ${total_spend:,.0f} USD (indicative, for planning purposes only)

Equipment inventory (exact quantities - do not modify):
{family_data}

CRITICAL LEAD TIME ITEMS (must order immediately):
{critical_str}

Recommended delivery sequence by floor:
{sequence_str}

TASKS:

1. RFQ LOTS - Group equipment into logical procurement batches based on supplier specialisation.
   CRITICAL RULES:
   (a) MUTUALLY EXCLUSIVE: assign each equipment family to ONE AND ONLY ONE lot.
   (b) EXACT QUANTITIES: list exact quantity per item as provided. Do not calculate totals or subtotals.
   (c) FLAG any items with critical lead times in their lot description.

2. 4-WEEK ACTION PLAN - Office-based procurement milestones only:
   - Week 1: Issue RFQs for CRITICAL lead time items immediately; shortlist suppliers
   - Week 2: Issue remaining RFQs; follow up on critical items
   - Week 3: Quotation analysis and negotiation
   - Week 4: Purchase order placement; confirm delivery schedule by floor level
   Note: if data quality issues exist, escalate Data Quality Log to BIM Manager. Do NOT suggest site visits.

3. TOP 3 RISKS - Prioritise lead time risk, spend concentration risk, and BIM data gaps.
   Include one mitigation per risk.

RULES: Do not recalculate quantities. Do not suggest site visits. No tables. Bullet points only. Professional tone.
"""


# ── EXCEL BUILDER (enhanced with 3 new sheets) ───────────────────────────────

def build_excel(df_clean, df_family, df_issues, df_lead, df_sequence, df_spend) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_clean.drop(columns=["Issues"], errors="ignore").to_excel(
            writer, sheet_name="Clean_BOM", index=False)
        df_family.to_excel(
            writer, sheet_name="Summary_by_Family", index=False)
        df_issues.to_excel(
            writer, sheet_name="Data_Quality_Log", index=False)
        df_lead.to_excel(
            writer, sheet_name="Lead_Time_Estimator", index=False)
        df_sequence.to_excel(
            writer, sheet_name="Delivery_Sequence")
        df_spend.to_excel(
            writer, sheet_name="Spend_Estimation", index=False)
    return output.getvalue()


# ── WORD BUILDER ──────────────────────────────────────────────────────────────

def build_word(df_clean, df_family, df_issues, total_units,
               ai_strategy, df_lead, df_sequence, df_spend, total_spend) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("Procurement Engineering Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"Total equipment units: {total_units}\n"
        f"Order lines (BOM rows): {len(df_clean)}\n"
        f"Data quality issues: {len(df_issues)}\n"
        f"Estimated procurement spend: ${total_spend:,.0f} USD"
    )

    doc.add_heading("2. AI-Generated Procurement Strategy", level=1)
    doc.add_paragraph(ai_strategy)

    doc.add_heading("3. Lead Time Estimator", level=1)
    doc.add_paragraph("Items sorted by order urgency. Dates based on project start date entered.")
    for _, row in df_lead.iterrows():
        doc.add_paragraph(
            f"{row['Status']} {row['Family']} — "
            f"Lead time: {row['Lead Time Min (wks)']}–{row['Lead Time Max (wks)']} wks — "
            f"Order by: {row['Order By Date']}",
            style="List Bullet"
        )

    doc.add_heading("4. Delivery Sequence by Floor", level=1)
    for _, row in df_sequence.iterrows():
        doc.add_paragraph(
            f"Phase {row.name}: {row['Level']} — {row['Total Units']} units — {row['Recommended Action']}",
            style="List Bullet"
        )

    doc.add_heading("5. Spend Estimation", level=1)
    doc.add_paragraph(f"Total estimated procurement value: ${total_spend:,.0f} USD (indicative)")
    for _, row in df_spend.iterrows():
        doc.add_paragraph(
            f"{row['Spend Tier']}  {row['Family']}: {row['Count']} units × "
            f"{row['Unit Cost (USD)']} = {row['Estimated Total (USD)']}",
            style="List Bullet"
        )

    doc.add_heading("6. Summary by Equipment Family", level=1)
    for _, row in df_family.iterrows():
        doc.add_paragraph(f"{row['Family']}: {row['Count']} units", style="List Bullet")

    doc.add_heading("7. Detailed BOM Line Items", level=1)
    for _, row in df_clean.iterrows():
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{row['Family']}").bold = True
        p.add_run(
            f" | Type: {row['Type']} | Panel: {row['Panel Name']} "
            f"| Level: {row['Level']} | Qty: {row['Count']}"
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTML BUILDER ──────────────────────────────────────────────────────────────

def build_html(df_clean, df_family, df_issues, total_units, ai_strategy,
               df_lead, df_sequence, df_spend, total_spend) -> str:

    family_list = "".join(
        f"<li><strong>{r['Family']}</strong>: {r['Count']} units</li>"
        for _, r in df_family.iterrows()
    )
    items_list = "".join(
        f"<div class='item-row'><strong>{r['Family']}</strong> — {r['Type']}<br>"
        f"<small>Panel: {r['Panel Name']} | Level: {r['Level']} | Qty: {r['Count']}</small></div>"
        for _, r in df_clean.iterrows()
    )
    issues_list = "".join(
        f"<li>{r['Family']} ({r['Type']}): {r['Issues']}</li>"
        for _, r in df_issues.iterrows()
    ) if len(df_issues) > 0 else "<li>No issues found.</li>"

    lead_rows = "".join(
        f"<tr><td>{r['Status']}</td><td>{r['Family']}</td>"
        f"<td>{r['Lead Time Min (wks)']}–{r['Lead Time Max (wks)']} wks</td>"
        f"<td>{r['Order By Date']}</td></tr>"
        for _, r in df_lead.iterrows()
    )
    seq_rows = "".join(
        f"<tr><td>{r.name}</td><td>{r['Level']}</td>"
        f"<td>{r['Total Units']}</td><td>{r['Recommended Action']}</td></tr>"
        for _, r in df_sequence.iterrows()
    )
    spend_rows = "".join(
        f"<tr><td>{r['Spend Tier']}</td><td>{r['Family']}</td>"
        f"<td>{r['Count']}</td><td>{r['Unit Cost (USD)']}</td>"
        f"<td>{r['Estimated Total (USD)']}</td></tr>"
        for _, r in df_spend.iterrows()
    )

    return f"""<!DOCTYPE html>
<html><head><meta charset='UTF-8'>
<title>Procurement Dashboard</title>
<style>
  body {{ font-family: 'Segoe UI', sans-serif; margin: 40px; background: #f8f9fa; line-height: 1.6; }}
  h1, h2 {{ color: #2c3e50; }}
  .card {{ background: white; padding: 25px; border-radius: 12px;
           box-shadow: 0 4px 10px rgba(0,0,0,0.05); margin-bottom: 25px; }}
  .stat-grid {{ display: flex; gap: 15px; margin-bottom: 20px; flex-wrap: wrap; }}
  .stat {{ flex: 1; min-width: 120px; padding: 20px; color: white; border-radius: 8px; text-align: center; }}
  .ai-box {{ border-left: 5px solid #f1c40f; padding-left: 15px;
             background: #fffdf0; white-space: pre-wrap; }}
  .item-row {{ padding: 10px; border-bottom: 1px solid #eee; font-size: 0.9rem; }}
  .item-row:hover {{ background: #f1f3f5; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 0.9rem; }}
  th {{ background: #2c3e50; color: white; padding: 8px 12px; text-align: left; }}
  td {{ padding: 8px 12px; border-bottom: 1px solid #eee; }}
  tr:hover {{ background: #f1f3f5; }}
</style>
</head>
<body>
<h1>⚡ Procurement Dashboard</h1>
<p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
<div class="stat-grid">
  <div class="stat" style="background:#3498db;"><h2>{total_units}</h2><p>Total Units</p></div>
  <div class="stat" style="background:#2ecc71;"><h2>{len(df_clean)}</h2><p>BOM Lines</p></div>
  <div class="stat" style="background:#e74c3c;"><h2>{len(df_issues)}</h2><p>Issues</p></div>
  <div class="stat" style="background:#8e44ad;"><h2>${total_spend/1000:.0f}k</h2><p>Est. Spend</p></div>
</div>

<div class="card" style="border-top:5px solid #f1c40f;">
  <h2>🤖 AI Procurement Strategy</h2>
  <div class="ai-box">{ai_strategy}</div>
</div>

<div class="card" style="border-top:5px solid #e74c3c;">
  <h2>⏱️ Lead Time Estimator</h2>
  <table><tr><th>Status</th><th>Family</th><th>Lead Time</th><th>Order By</th></tr>
  {lead_rows}</table>
</div>

<div class="card" style="border-top:5px solid #3498db;">
  <h2>🚚 Delivery Sequence by Floor</h2>
  <table><tr><th>Phase</th><th>Level</th><th>Units</th><th>Recommended Action</th></tr>
  {seq_rows}</table>
</div>

<div class="card" style="border-top:5px solid #8e44ad;">
  <h2>💰 Spend Estimation (Indicative)</h2>
  <p><strong>Total estimated procurement value: ${total_spend:,.0f} USD</strong></p>
  <table><tr><th>Tier</th><th>Family</th><th>Qty</th><th>Unit Cost</th><th>Total</th></tr>
  {spend_rows}</table>
</div>

<div class="card">
  <h2>📦 Summary by Family</h2>
  <ul>{family_list}</ul>
</div>

<div class="card">
  <h2>⚠️ Data Quality Issues</h2>
  <ul>{issues_list}</ul>
</div>

<div class="card">
  <h2>📋 Detailed BOM ({len(df_clean)} lines)</h2>
  <div style="max-height:400px;overflow-y:auto;">{items_list}</div>
</div>
</body></html>"""


# ══════════════════════════════════════════════════════════════════════════════
# MAIN APP
# ══════════════════════════════════════════════════════════════════════════════

st.title("⚡ BIM Procurement Assistant")
st.caption("TFM Tool — BIM Schedule → Procurement Pipeline + AI Strategy + Analytics")

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Configuration")

    st.subheader("🤖 AI Provider")
    ai_provider = st.selectbox(
        "Select AI",
        ["Groq (free)", "Ollama (local)", "OpenAI", "Claude (Anthropic)"],
        help="Groq is free and works in the cloud. Ollama runs locally."
    )

    api_key = ""
    model_name = ""

    if ai_provider == "Groq (free)":
        api_key = st.text_input("Groq API Key", type="password",
                                help="Free at console.groq.com → API Keys")
        model_name = st.selectbox("Model", [
            "llama-3.1-8b-instant",
            "llama-3.3-70b-versatile",
            "mixtral-8x7b-32768",
        ])
        st.success("✅ Groq is free — no credit card needed!")

    elif ai_provider == "Ollama (local)":
        model_name = st.text_input("Ollama model", value="llama3")

    elif ai_provider == "OpenAI":
        api_key = st.text_input("OpenAI API Key", type="password")
        model_name = st.selectbox("Model", ["gpt-4o-mini", "gpt-4o", "gpt-3.5-turbo"])

    elif ai_provider == "Claude (Anthropic)":
        api_key = st.text_input("Anthropic API Key", type="password")
        model_name = st.selectbox("Model", [
            "claude-haiku-4-5-20251001",
            "claude-sonnet-4-6",
        ])

    st.divider()
    st.subheader("📅 Project Settings")
    project_start = st.date_input(
        "Target Installation Start Date",
        value=datetime.now() + timedelta(weeks=20),
        help="The pipeline calculates order deadlines working backwards from this date."
    )
    project_start_dt = datetime.combine(project_start, datetime.min.time())

    st.divider()
    st.info(
        "**How to use:**\n"
        "1. Set your installation start date\n"
        "2. Upload your Revit CSV\n"
        "3. Click 'Run Pipeline'\n"
        "4. Review results & download"
    )

# ── MAIN AREA ─────────────────────────────────────────────────────────────────
col_upload, col_run = st.columns([3, 1])

with col_upload:
    uploaded_file = st.file_uploader(
        "📂 Upload Revit Equipment Schedule (CSV)",
        type=["csv"],
        help="Export directly from Revit without any manual editing."
    )

with col_run:
    st.write("")
    st.write("")
    run_btn = st.button("🚀 Run Pipeline", type="primary", use_container_width=True,
                        disabled=(uploaded_file is None))

# ── PIPELINE ──────────────────────────────────────────────────────────────────
if run_btn and uploaded_file:

    with st.spinner("Step 1/4 — Cleaning BIM data..."):
        df_clean, df_family, df_issues, total_units = clean_bim_csv(uploaded_file)

    with st.spinner("Step 2/4 — Running procurement analytics..."):
        df_lead     = build_lead_time_df(df_family, project_start_dt)
        df_sequence = build_delivery_sequence(df_clean)
        df_spend    = build_spend_estimation(df_family)
        total_spend = get_total_spend(df_family)

    # ── KPI METRICS ──────────────────────────────────────────────────────────
    st.divider()
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("⚡ Total Units", total_units)
    c2.metric("📋 BOM Lines", len(df_clean))
    c3.metric("⚠️ Issues", len(df_issues))
    c4.metric("📦 Families", len(df_family))
    c5.metric("💰 Est. Spend", f"${total_spend/1000:.0f}k")

    # ── PREVIEW TABS ─────────────────────────────────────────────────────────
    st.subheader("📊 Data Preview")
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "✅ Clean BOM", "📦 By Family", "⚠️ Quality Issues",
        "⏱️ Lead Times", "🚚 Delivery Sequence", "💰 Spend Estimation"
    ])

    with tab1:
        st.dataframe(df_clean.drop(columns=["Issues"], errors="ignore"),
                     use_container_width=True, height=300)

    with tab2:
        st.bar_chart(df_family.set_index("Family")["Count"])
        st.dataframe(df_family, use_container_width=True)

    with tab3:
        if len(df_issues) > 0:
            st.warning(f"{len(df_issues)} rows with missing procurement attributes.")
            st.dataframe(df_issues[["Family", "Type", "Panel Name", "Level", "Count", "Issues"]],
                         use_container_width=True)
        else:
            st.success("No data quality issues found!")

    with tab4:
        st.info(f"Order deadlines calculated backwards from installation start: **{project_start}**")
        critical_count = len(df_lead[df_lead["Status"].str.contains("OVERDUE|CRITICAL", na=False)])
        if critical_count > 0:
            st.error(f"⚠️ {critical_count} item(s) require immediate ordering action!")
        st.dataframe(df_lead, use_container_width=True, height=350)

    with tab5:
        st.info("Recommended delivery phasing based on floor level distribution.")
        st.dataframe(df_sequence, use_container_width=True)

    with tab6:
        st.info(f"**Total estimated procurement value: ${total_spend:,.0f} USD** (indicative — based on published industry benchmarks)")
        st.dataframe(df_spend, use_container_width=True, height=350)
        st.caption("Unit costs sourced from RS Means / NECA Manual of Labor Units benchmarks. For budgeting reference only.")

    # ── AI STRATEGY ──────────────────────────────────────────────────────────
    st.divider()
    st.subheader("🤖 AI Procurement Strategy")

    with st.spinner(f"Step 3/4 — Calling {ai_provider}..."):
        prompt = build_prompt(
            total_units, len(df_clean), len(df_issues),
            df_family, df_lead, df_sequence, total_spend
        )
        ai_strategy = call_ai(ai_provider, api_key, model_name, prompt)

    st.markdown(f"<div class='ai-box'>{ai_strategy}</div>", unsafe_allow_html=True)

    # ── EXPORTS ──────────────────────────────────────────────────────────────
    st.divider()
    st.subheader("📥 Download Outputs")

    with st.spinner("Step 4/4 — Building export files..."):
        excel_bytes = build_excel(df_clean, df_family, df_issues,
                                  df_lead, df_sequence, df_spend)
        word_bytes  = build_word(df_clean, df_family, df_issues, total_units,
                                 ai_strategy, df_lead, df_sequence, df_spend, total_spend)
        html_str    = build_html(df_clean, df_family, df_issues, total_units,
                                 ai_strategy, df_lead, df_sequence, df_spend, total_spend)

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w") as zf:
            zf.writestr("procurement_outputs/tfm_outputs_procurement.xlsx", excel_bytes)
            zf.writestr("procurement_outputs/Procurement_Technical_Annex.docx", word_bytes)
            zf.writestr("procurement_outputs/Procurement_Report.html", html_str.encode("utf-8"))
        zip_bytes = zip_buf.getvalue()

    dl1, dl2, dl3, dl4 = st.columns(4)
    with dl1:
        st.download_button("📊 Download Excel", data=excel_bytes,
            file_name="tfm_outputs_procurement.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True)
    with dl2:
        st.download_button("📝 Download Word", data=word_bytes,
            file_name="Procurement_Technical_Annex.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True)
    with dl3:
        st.download_button("🌐 Download HTML", data=html_str.encode("utf-8"),
            file_name="Procurement_Report.html", mime="text/html",
            use_container_width=True)
    with dl4:
        st.download_button("📦 Download All (ZIP)", data=zip_bytes,
            file_name="procurement_outputs.zip", mime="application/zip",
            use_container_width=True)

    st.success("✅ Pipeline complete! All files ready to download.")

elif not uploaded_file:
    st.info("👆 Upload a Revit CSV file to get started.")


# ── AI CALL (unchanged from v1) ───────────────────────────────────────────────
def call_ai(ai_provider, api_key, model_name, prompt) -> str:
    if ai_provider == "Ollama (local)":
        try:
            import ollama
            response = ollama.chat(
                model=model_name or "llama3",
                messages=[{"role": "user", "content": prompt}]
            )
            return response["message"]["content"]
        except Exception as e:
            return f"[Ollama error] {e}"

    elif ai_provider == "Groq (free)":
        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key, base_url="https://api.groq.com/openai/v1")
            response = client.chat.completions.create(
                model=model_name or "llama-3.1-8b-instant",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"[Groq error] {e}"

    elif ai_provider == "OpenAI":
        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model=model_name or "gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"[OpenAI error] {e}"

    elif ai_provider == "Claude (Anthropic)":
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            message = client.messages.create(
                model=model_name or "claude-haiku-4-5-20251001",
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt}],
            )
            return message.content[0].text
        except Exception as e:
            return f"[Anthropic error] {e}"

    return "No AI provider configured."
